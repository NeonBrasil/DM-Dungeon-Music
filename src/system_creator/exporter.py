"""
DM - Dungeon Music
Exportador de sistemas e fichas de personagem.
Suporta: TXT, Word (.docx), PNG, PDF
  — Documento do Sistema: TXT, Word, PNG, PDF
  — Ficha de Personagem:  TXT, Word, PNG, PDF
"""

import os
import textwrap
from datetime import datetime

try:
    from PIL import Image, ImageDraw, ImageFont
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    )
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

_EXPORT_DIR = os.path.join(os.path.expanduser("~"), ".dm_dungeon_music", "exports")


def _ensure_dir():
    os.makedirs(_EXPORT_DIR, exist_ok=True)


def _safe_filename(name: str) -> str:
    return "".join(c if c.isalnum() or c in " _-" else "_" for c in name).strip()


def _fmt_list(value) -> str:
    if isinstance(value, list):
        return ", ".join(str(v) for v in value)
    return str(value) if value else ""


# ─────────────────────────────────────────────────────────────────────────────
# HELPERS DE CONTEÚDO
# ─────────────────────────────────────────────────────────────────────────────

def _system_sections(system_data: dict) -> list:
    """Retorna lista de (título, [(label, valor)]) para o documento do sistema."""
    answers = system_data.get("answers", {})
    name = system_data.get("name", "Sistema")
    sections = []

    sections.append(("Fundação", [
        ("Nome do Sistema",   answers.get("system_name", name)),
        ("Sistema Base",      answers.get("base_system", "")),
        ("Gênero / Tom",      _fmt_list(answers.get("genre", []))),
    ]))

    sections.append(("Mecânicas", [
        ("Resolução de Ações",    answers.get("core_resolution", "")),
        ("Dados utilizados",      _fmt_list(answers.get("dice_types", []))),
        ("Resultados",            answers.get("success_failure", "")),
        ("Sistema de Saúde",      answers.get("health_system", "")),
        ("Recuperação",           _fmt_list(answers.get("recovery_rest", []))),
        ("Magia / Poderes",       _fmt_list(answers.get("magic_system", []))),
        ("Progressão",            answers.get("advancement", "")),
        ("Estilo de Combate",     _fmt_list(answers.get("combat_style", []))),
        ("Interação Social",      answers.get("social_mechanics", "")),
        ("Categorias de Itens",   _fmt_list(answers.get("equipment_categories", []))),
    ]))

    sections.append(("Personagens", [
        ("Atributos",             answers.get("attributes_type", "")),
        ("Atributos custom",      _fmt_list(answers.get("custom_attributes", []))),
        ("Perícias",              answers.get("skills_system", "")),
        ("Arquétipos / Classes",  answers.get("archetypes", "")),
        ("Vantagens & Falhas",    _fmt_list(answers.get("special_traits", []))),
        ("Habilidades Chave",     _fmt_list(answers.get("key_abilities", []))),
        ("Morte & Consequências", answers.get("death_rules", "")),
    ]))

    if answers.get("setting"):
        sections.append(("Mundo & Lore", [
            ("Cenário",          answers.get("setting", "")),
            ("Facções",          _fmt_list(answers.get("factions", []))),
            ("Temas Centrais",   _fmt_list(answers.get("themes", []))),
        ]))

    sheet_cfg = system_data.get("sheet_config", {})
    fields = sheet_cfg.get("fields", [])
    sections.append(("Ficha de Personagem", [
        ("Layout",           answers.get("sheet_layout", "")),
        ("Campos",           ", ".join(f.get("label", "") for f in fields)),
    ]))

    if answers.get("final_notes"):
        sections.append(("Notas Finais", [
            ("Notas", answers.get("final_notes", "")),
        ]))

    return sections


def _sheet_rows(system_data: dict) -> list:
    """Retorna lista de (label, tipo) para a ficha."""
    fields = system_data.get("sheet_config", {}).get("fields", [])
    return [(f.get("label", ""), f.get("type", "")) for f in fields]


# ─────────────────────────────────────────────────────────────────────────────
# TXT — DOCUMENTO DO SISTEMA
# ─────────────────────────────────────────────────────────────────────────────

def export_system_txt(system_data: dict, output_path: str) -> str:
    name = system_data.get("name", "Sistema")
    sep = "=" * 72
    thin = "-" * 72
    lines = []

    lines += [sep, f"  SISTEMA: {name.upper()}", sep,
              f"  Criado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}", ""]

    for section_title, items in _system_sections(system_data):
        lines += ["", f"── {section_title} ──", thin]
        for label, value in items:
            if not value:
                continue
            wrapped = textwrap.wrap(str(value), 64)
            if len(wrapped) <= 1:
                lines.append(f"  {label}: {value}")
            else:
                lines.append(f"  {label}:")
                for w in wrapped:
                    lines.append(f"    {w}")

    lines += ["", sep, "  Gerado por DM - Dungeon Music", sep]
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return output_path


# ─────────────────────────────────────────────────────────────────────────────
# TXT — FICHA DE PERSONAGEM
# ─────────────────────────────────────────────────────────────────────────────

def export_sheet_txt(system_data: dict, output_path: str) -> str:
    name = system_data.get("name", "Sistema")
    answers = system_data.get("answers", {})
    sep = "=" * 72
    thin = "-" * 72
    lines = []

    lines += [sep, f"  FICHA DE PERSONAGEM — {name.upper()}", sep,
              f"  Sistema Base: {answers.get('base_system', '')}",
              f"  Mestre: ________________________  Data: {datetime.now().strftime('%d/%m/%Y')}", ""]

    for label, ftype in _sheet_rows(system_data):
        lines.append(thin)
        if ftype == "track":
            lines.append(f"  {label}: □□□□□□□□□□")
        elif ftype == "dots":
            lines.append(f"  {label}: ○○○○○")
        elif ftype == "number":
            lines.append(f"  {label}: _______")
        elif ftype in ("text_long", "text"):
            lines.append(f"  {label}:")
            lines.append("  " + "_" * 60)
        elif ftype == "table":
            lines.append(f"  {label}:")
            lines.append("  " + "-" * 48)
            for _ in range(3):
                lines.append("  |" + " " * 46 + "|")
            lines.append("  " + "-" * 48)
        elif ftype == "list":
            lines.append(f"  {label}:")
            for _ in range(4):
                lines.append("    • " + "_" * 40)
        elif ftype == "checkbox":
            lines.append(f"  {label}: □ Sim  □ Não")
        else:
            lines.append(f"  {label}: _______")
        lines.append("")

    lines += [sep, "  Gerado por DM - Dungeon Music", sep]
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return output_path


# ─────────────────────────────────────────────────────────────────────────────
# WORD — DOCUMENTO DO SISTEMA
# ─────────────────────────────────────────────────────────────────────────────

def export_system_word(system_data: dict, output_path: str) -> str:
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx não instalado. Use: pip install python-docx")

    name = system_data.get("name", "Sistema")
    doc = DocxDocument()

    # Margens
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Título
    h = doc.add_heading(f"Sistema: {name}", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    answers = system_data.get("answers", {})
    p = doc.add_paragraph()
    p.add_run(f"Base: {answers.get('base_system', 'Personalizado')}  •  "
              f"Criado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for section_title, items in _system_sections(system_data):
        doc.add_heading(section_title, level=1)
        for label, value in items:
            if not value:
                continue
            p = doc.add_paragraph(style="List Bullet")
            run_lbl = p.add_run(f"{label}: ")
            run_lbl.bold = True
            p.add_run(str(value))
        doc.add_paragraph()

    doc.add_paragraph("─" * 50)
    footer_p = doc.add_paragraph("Gerado por DM - Dungeon Music")
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)
    return output_path


# ─────────────────────────────────────────────────────────────────────────────
# WORD — FICHA DE PERSONAGEM
# ─────────────────────────────────────────────────────────────────────────────

def export_sheet_word(system_data: dict, output_path: str) -> str:
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx não instalado. Use: pip install python-docx")

    name = system_data.get("name", "Sistema")
    answers = system_data.get("answers", {})
    doc = DocxDocument()

    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    # Cabeçalho
    h = doc.add_heading("FICHA DE PERSONAGEM", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f"Sistema: {name}  •  Base: {answers.get('base_system', '')}  •  "
                           f"Mestre: _________________")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Campo Nome em destaque
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    cell.width = Cm(16)
    run = cell.paragraphs[0].add_run("Nome do Personagem: ")
    run.bold = True
    cell.paragraphs[0].add_run("_" * 40)
    doc.add_paragraph()

    # Campos em 2 colunas (tabela 2 cols)
    rows_data = _sheet_rows(system_data)
    i = 0
    while i < len(rows_data):
        label, ftype = rows_data[i]
        width_2 = ftype in ("text_long", "table", "list", "image")

        if width_2:
            tbl = doc.add_table(rows=1, cols=1)
            tbl.style = "Table Grid"
            cell = tbl.cell(0, 0)
            _fill_sheet_cell_word(cell, label, ftype)
            i += 1
        else:
            # Pega par de campos para 2 colunas
            tbl = doc.add_table(rows=1, cols=2)
            tbl.style = "Table Grid"
            _fill_sheet_cell_word(tbl.cell(0, 0), label, ftype)
            if i + 1 < len(rows_data) and rows_data[i + 1][1] not in ("text_long", "table", "list", "image"):
                _fill_sheet_cell_word(tbl.cell(0, 1), rows_data[i + 1][0], rows_data[i + 1][1])
                i += 2
            else:
                tbl.cell(0, 1).text = ""
                i += 1
        doc.add_paragraph()

    doc.add_paragraph("─" * 50)
    footer = doc.add_paragraph("Gerado por DM - Dungeon Music")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)
    return output_path


def _fill_sheet_cell_word(cell, label: str, ftype: str):
    """Preenche uma célula da tabela Word com o campo da ficha."""
    p = cell.paragraphs[0]
    run = p.add_run(f"{label.upper()}\n")
    run.bold = True
    run.font.size = Pt(9)

    if ftype == "track":
        p.add_run("□ " * 10)
    elif ftype == "dots":
        p.add_run("○ " * 5)
    elif ftype == "number":
        p.add_run("______")
    elif ftype in ("text", "text_long"):
        lines = 4 if ftype == "text_long" else 1
        for _ in range(lines):
            p.add_run("\n" + "_" * 35)
    elif ftype == "table":
        p.add_run("[ tabela ]")
    elif ftype == "list":
        for _ in range(4):
            p.add_run("\n• " + "_" * 30)
    elif ftype == "checkbox":
        p.add_run("□ Sim   □ Não")
    elif ftype == "image":
        p.add_run("[ Retrato / Imagem ]")
    else:
        p.add_run("______")


# ─────────────────────────────────────────────────────────────────────────────
# PNG — DOCUMENTO DO SISTEMA (texto renderizado)
# ─────────────────────────────────────────────────────────────────────────────

_DARK = {
    "bg":      (18, 18, 28),
    "surface": (28, 28, 42),
    "card":    (38, 38, 56),
    "border":  (80, 80, 110),
    "primary": (99, 102, 241),
    "text":    (230, 230, 240),
    "muted":   (140, 140, 160),
    "accent":  (167, 139, 250),
}

_W = 1240
_H = 1754


def _load_font(size: int, bold: bool = False):
    if not PIL_AVAILABLE:
        return None
    candidates = (
        ["C:/Windows/Fonts/calibrib.ttf", "C:/Windows/Fonts/arialbd.ttf",
         "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf"]
        if bold else
        ["C:/Windows/Fonts/calibri.ttf", "C:/Windows/Fonts/arial.ttf",
         "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"]
    )
    for path in candidates:
        if os.path.isfile(path):
            try:
                return ImageFont.truetype(path, size)
            except Exception:
                pass
    try:
        return ImageFont.load_default()
    except Exception:
        return None


def export_system_png(system_data: dict, output_path: str) -> str:
    if not PIL_AVAILABLE:
        raise RuntimeError("Pillow não instalado. Use: pip install Pillow")

    name = system_data.get("name", "Sistema")
    answers = system_data.get("answers", {})

    img = Image.new("RGB", (_W, _H), _DARK["bg"])
    draw = ImageDraw.Draw(img)

    font_title  = _load_font(36, bold=True)
    font_h2     = _load_font(20, bold=True)
    font_label  = _load_font(14, bold=True)
    font_normal = _load_font(14)

    # Header
    draw.rectangle([0, 0, _W, 80], fill=_DARK["card"])
    draw.line([0, 80, _W, 80], fill=_DARK["primary"], width=3)
    draw.text((40, 14), f"SISTEMA: {name.upper()}", font=font_title, fill=_DARK["primary"])
    draw.text((40, 56), f"Base: {answers.get('base_system', '')}  •  "
              f"{datetime.now().strftime('%d/%m/%Y')}", font=font_normal, fill=_DARK["muted"])

    y = 100
    pad = 40
    line_h = 22

    for section_title, items in _system_sections(system_data):
        if y > _H - 120:
            break
        # Título da seção
        draw.rectangle([pad, y, _W - pad, y + 32], fill=_DARK["surface"])
        draw.text((pad + 10, y + 6), section_title.upper(), font=font_h2, fill=_DARK["accent"])
        y += 38

        for label, value in items:
            if not value or y > _H - 80:
                continue
            text_full = f"{label}: {value}"
            wrapped = textwrap.wrap(text_full, 90)
            for j, line in enumerate(wrapped):
                fnt = font_label if j == 0 else font_normal
                col = _DARK["text"] if j == 0 else _DARK["muted"]
                draw.text((pad + 16, y), line, font=fnt, fill=col)
                y += line_h
        y += 8

    # Footer
    draw.rectangle([0, _H - 40, _W, _H], fill=_DARK["card"])
    draw.line([0, _H - 40, _W, _H - 40], fill=_DARK["border"], width=1)
    draw.text((40, _H - 28), "DM - Dungeon Music", font=font_normal, fill=_DARK["muted"])

    img.save(output_path, "PNG")
    return output_path


# ─────────────────────────────────────────────────────────────────────────────
# PNG — FICHA DE PERSONAGEM (visual)
# ─────────────────────────────────────────────────────────────────────────────

def export_sheet_png(system_data: dict, output_path: str) -> str:
    if not PIL_AVAILABLE:
        raise RuntimeError("Pillow não instalado. Use: pip install Pillow")

    name = system_data.get("name", "Sistema")
    answers = system_data.get("answers", {})
    fields = system_data.get("sheet_config", {}).get("fields", [])

    img = Image.new("RGB", (_W, _H), _DARK["bg"])
    draw = ImageDraw.Draw(img)

    font_title  = _load_font(36, bold=True)
    font_h2     = _load_font(20, bold=True)
    font_label  = _load_font(14, bold=True)
    font_normal = _load_font(14)
    font_small  = _load_font(12)

    # Header
    draw.rectangle([0, 0, _W, 80], fill=_DARK["card"])
    draw.line([0, 80, _W, 80], fill=_DARK["primary"], width=3)
    draw.text((40, 14), "FICHA DE PERSONAGEM", font=font_title, fill=_DARK["primary"])
    draw.text((40, 56), f"Sistema: {name}  •  Base: {answers.get('base_system', '')}",
              font=font_normal, fill=_DARK["muted"])

    y = 100
    pad = 30
    col_w = (_W - pad * 2 - 20) // 2
    col_x = [pad, pad + col_w + 20]
    col_y = [y, y]
    row_h_base = 68

    def draw_field(field_dict, col):
        ftype  = field_dict.get("type", "text")
        flabel = field_dict.get("label", "")
        fwidth = field_dict.get("width", 1)
        cx = col_x[col]
        cy = col_y[col]
        fw = col_w if fwidth == 1 else (_W - pad * 2)
        fh = row_h_base

        if ftype == "text_long": fh = 100
        elif ftype == "table":   fh = 110
        elif ftype == "list":    fh = 110
        elif ftype == "image":   fh = 160

        if cy + fh > _H - 50:
            return

        draw.rectangle([cx, cy, cx + fw, cy + fh],
                        fill=_DARK["surface"], outline=_DARK["border"], width=1)
        draw.text((cx + 8, cy + 5), flabel.upper(), font=font_label, fill=_DARK["muted"])
        draw.line([cx + 8, cy + 24, cx + fw - 8, cy + 24], fill=_DARK["border"])

        iy = cy + 30
        if ftype == "track":
            maxi = min(field_dict.get("max", 10), 15)
            for k in range(maxi):
                bx = cx + 10 + k * 18
                draw.rectangle([bx, iy, bx + 14, iy + 14], outline=_DARK["border"])
        elif ftype == "dots":
            maxi = field_dict.get("max", 5)
            for k in range(maxi):
                bx = cx + 10 + k * 22
                draw.ellipse([bx, iy, bx + 14, iy + 14], outline=_DARK["muted"])
        elif ftype == "table":
            cols_t = field_dict.get("columns", ["Campo", "Valor"])
            cw2 = (fw - 16) // len(cols_t)
            for ci, ch in enumerate(cols_t):
                hx = cx + 8 + ci * cw2
                draw.rectangle([hx, iy, hx + cw2 - 2, iy + 18], fill=_DARK["card"])
                draw.text((hx + 3, iy + 2), ch, font=font_small, fill=_DARK["accent"])
            for ri in range(3):
                ry = iy + 20 + ri * 22
                for ci in range(len(cols_t)):
                    hx = cx + 8 + ci * cw2
                    draw.rectangle([hx, ry, hx + cw2 - 2, ry + 20], outline=_DARK["border"])
        elif ftype in ("text", "text_long"):
            lines_n = 1 if ftype == "text" else 3
            for li in range(lines_n):
                ly = iy + li * 24
                draw.line([cx + 8, ly + 18, cx + fw - 8, ly + 18], fill=_DARK["border"])
        elif ftype == "list":
            for li in range(4):
                ly = iy + li * 22
                draw.text((cx + 8, ly), "•", font=font_normal, fill=_DARK["muted"])
                draw.line([cx + 24, ly + 16, cx + fw - 8, ly + 16], fill=_DARK["border"])
        elif ftype == "image":
            draw.rectangle([cx + 8, iy, cx + fw - 8, cy + fh - 8],
                            fill=_DARK["card"], outline=_DARK["border"])
            draw.text((cx + fw // 2 - 30, iy + 55), "[ Retrato ]",
                      font=font_normal, fill=_DARK["muted"])

        step = fh + 8
        if fwidth == 2:
            col_y[0] += step
            col_y[1] = col_y[0]
        else:
            col_y[col] += step

    current_col = 0
    for field in fields:
        if field.get("width", 1) == 2:
            current_col = 0
        if col_y[current_col] > _H - 80:
            break
        draw_field(field, current_col)
        if field.get("width", 1) == 1:
            current_col = 1 - current_col

    # Footer
    draw.rectangle([0, _H - 40, _W, _H], fill=_DARK["card"])
    draw.line([0, _H - 40, _W, _H - 40], fill=_DARK["border"])
    draw.text((40, _H - 28),
              f"DM - Dungeon Music  •  Mestre: _________________________  {datetime.now().strftime('%d/%m/%Y')}",
              font=font_small, fill=_DARK["muted"])

    img.save(output_path, "PNG")
    return output_path


# ─────────────────────────────────────────────────────────────────────────────
# PDF — DOCUMENTO DO SISTEMA
# ─────────────────────────────────────────────────────────────────────────────

def export_system_pdf(system_data: dict, output_path: str) -> str:
    if not REPORTLAB_AVAILABLE:
        raise RuntimeError("reportlab não instalado. Use: pip install reportlab")

    name = system_data.get("name", "Sistema")
    answers = system_data.get("answers", {})

    doc = SimpleDocTemplate(output_path, pagesize=A4,
                             rightMargin=2*cm, leftMargin=2*cm,
                             topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()

    title_s = ParagraphStyle("T", parent=styles["Title"], fontSize=22,
                               textColor=colors.HexColor("#6366f1"), spaceAfter=4)
    h2_s    = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=13,
                               textColor=colors.HexColor("#a78bfa"), spaceBefore=10, spaceAfter=4)
    body_s  = ParagraphStyle("B", parent=styles["Normal"], fontSize=10,
                               textColor=colors.HexColor("#1a1a2e"), spaceAfter=3)
    lbl_s   = ParagraphStyle("L", parent=body_s, fontName="Helvetica-Bold")

    story = []
    story.append(Paragraph(f"Sistema: {name}", title_s))
    story.append(Paragraph(
        f"Base: {answers.get('base_system', '')}  •  "
        f"Criado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}", body_s))
    story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor("#6366f1")))
    story.append(Spacer(1, 0.3*cm))

    for section_title, items in _system_sections(system_data):
        story.append(Paragraph(section_title, h2_s))
        for label, value in items:
            if value:
                story.append(Paragraph(f"<b>{label}:</b> {value}", body_s))
        story.append(Spacer(1, 0.2*cm))

    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#cccccc")))
    story.append(Paragraph("Gerado por DM - Dungeon Music", body_s))
    doc.build(story)
    return output_path


# ─────────────────────────────────────────────────────────────────────────────
# PDF — FICHA DE PERSONAGEM
# ─────────────────────────────────────────────────────────────────────────────

def export_sheet_pdf(system_data: dict, output_path: str) -> str:
    if not REPORTLAB_AVAILABLE:
        raise RuntimeError("reportlab não instalado. Use: pip install reportlab")

    from reportlab.platypus import KeepTogether
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    name    = system_data.get("name", "Sistema")
    answers = system_data.get("answers", {})
    fields  = system_data.get("sheet_config", {}).get("fields", [])
    base    = answers.get("base_system", "")

    # ── Cores ─────────────────────────────────────────────────────────────────
    C_PRIMARY  = colors.HexColor("#5b21b6")   # roxo escuro
    C_ACCENT   = colors.HexColor("#7c3aed")   # roxo médio
    C_LIGHT    = colors.HexColor("#ede9fe")   # roxo muito claro
    C_HEADER   = colors.HexColor("#f5f3ff")   # fundo seção
    C_BORDER   = colors.HexColor("#c4b5fd")   # borda
    C_MUTED    = colors.HexColor("#6b7280")   # texto muted
    C_TEXT     = colors.HexColor("#1e1b4b")   # texto escuro
    C_WHITE    = colors.white
    C_TRACK_BG = colors.HexColor("#ddd6fe")   # caixa track preenchida
    C_TRACK_EM = colors.HexColor("#f3f0ff")   # caixa track vazia

    # ── Estilos ────────────────────────────────────────────────────────────────
    styles = getSampleStyleSheet()
    title_s  = ParagraphStyle("ShT",  fontName="Helvetica-Bold", fontSize=22,
                               textColor=C_WHITE, alignment=TA_CENTER, leading=26)
    sub_s    = ParagraphStyle("ShS",  fontName="Helvetica",      fontSize=8,
                               textColor=C_WHITE, alignment=TA_CENTER, spaceAfter=0)
    sec_s    = ParagraphStyle("ShSc", fontName="Helvetica-Bold", fontSize=9,
                               textColor=C_ACCENT, spaceAfter=2, spaceBefore=2)
    lbl_s    = ParagraphStyle("ShL",  fontName="Helvetica-Bold", fontSize=7,
                               textColor=C_MUTED, spaceAfter=1)
    body_s   = ParagraphStyle("ShB",  fontName="Helvetica",      fontSize=9,
                               textColor=C_TEXT,  spaceAfter=1)
    track_s  = ParagraphStyle("ShTr", fontName="Helvetica",      fontSize=11,
                               textColor=C_ACCENT, spaceAfter=0, leading=14)
    blank_s  = ParagraphStyle("ShBl", fontName="Helvetica",      fontSize=9,
                               textColor=C_MUTED,  spaceAfter=2)

    doc = SimpleDocTemplate(
        output_path, pagesize=A4,
        rightMargin=1.2*cm, leftMargin=1.2*cm,
        topMargin=0.8*cm,   bottomMargin=1*cm,
    )

    story = []
    FULL_W = 17.6 * cm
    COL_W  = [FULL_W / 2 - 0.15*cm, FULL_W / 2 - 0.15*cm]

    # ── Banner de título ───────────────────────────────────────────────────────
    banner_data = [[
        Paragraph("FICHA DE PERSONAGEM", title_s),
    ]]
    banner = Table(banner_data, colWidths=[FULL_W])
    banner.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), C_PRIMARY),
        ("PADDING",    (0, 0), (-1, -1), 10),
        ("ROUNDEDCORNERS", [4]),
    ]))
    story.append(banner)
    story.append(Spacer(1, 0.2*cm))

    # Sub-linha de meta
    meta_data = [[
        Paragraph(f"Sistema: <b>{name}</b>   |   Base: {base}   |   "
                  f"Data: {datetime.now().strftime('%d/%m/%Y')}", sub_s),
    ]]
    meta_tbl = Table(meta_data, colWidths=[FULL_W])
    meta_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), C_ACCENT),
        ("PADDING",    (0, 0), (-1, -1), 5),
    ]))
    story.append(meta_tbl)
    story.append(Spacer(1, 0.3*cm))

    # ── Campo Nome + Jogador em destaque ──────────────────────────────────────
    identity_data = [
        [Paragraph("NOME DO PERSONAGEM", lbl_s),
         Paragraph("JOGADOR", lbl_s)],
        [Paragraph("_" * 45, blank_s),
         Paragraph("_" * 30, blank_s)],
    ]
    identity_tbl = Table(identity_data, colWidths=[FULL_W * 0.6, FULL_W * 0.4])
    identity_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0),  C_LIGHT),
        ("BACKGROUND", (0, 1), (-1, 1),  C_WHITE),
        ("BOX",        (0, 0), (-1, -1), 1, C_BORDER),
        ("INNERGRID",  (0, 0), (-1, -1), 0.5, C_BORDER),
        ("PADDING",    (0, 0), (-1, -1), 6),
        ("VALIGN",     (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.append(identity_tbl)
    story.append(Spacer(1, 0.25*cm))

    # ── Helpers de célula ─────────────────────────────────────────────────────

    def _track_boxes(count: int, label_list: list = None) -> list:
        """Linha de caixas de trilha."""
        box = "■ " * 3 + "□ " * max(0, count - 3)   # primeiras 3 preenchidas como modelo
        items = [Paragraph(box.strip(), track_s)]
        if label_list:
            items.append(Paragraph(" / ".join(label_list), blank_s))
        return items

    def _make_cell(label: str, ftype: str, field_dict: dict = None) -> list:
        fd = field_dict or {}
        items = [Paragraph(label.upper(), lbl_s)]

        if ftype == "track":
            maxi  = fd.get("max", 10)
            lbls  = fd.get("labels", [])
            boxes = ("■ " * 3 + "□ " * max(0, maxi - 3)).strip()
            items.append(Paragraph(boxes, track_s))
            if lbls:
                items.append(Paragraph("  /  ".join(lbls), blank_s))

        elif ftype == "dots":
            maxi  = fd.get("max", 5)
            items.append(Paragraph("● ● ●   ○  ○  " if maxi >= 5 else "● " * maxi, track_s))

        elif ftype == "number":
            items.append(Paragraph("________", blank_s))

        elif ftype == "text":
            items.append(Paragraph("_" * 35, blank_s))

        elif ftype == "text_long":
            for _ in range(4):
                items.append(Paragraph("_" * 35, blank_s))

        elif ftype == "table":
            cols_t = fd.get("columns", ["Campo", "Valor"])
            # Mini tabela de linhas
            header = [Paragraph(c, lbl_s) for c in cols_t]
            trows  = [header]
            for _ in range(3):
                trows.append([Paragraph("", blank_s)] * len(cols_t))
            inner = Table(trows)
            inner.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0),  C_LIGHT),
                ("GRID",       (0, 0), (-1, -1), 0.4, C_BORDER),
                ("PADDING",    (0, 0), (-1, -1), 3),
                ("FONTSIZE",   (0, 0), (-1, -1), 7),
            ]))
            items.append(inner)

        elif ftype == "list":
            for j in range(5):
                items.append(Paragraph(f"{'①②③④⑤'[j]}  " + "_" * 28, blank_s))

        elif ftype == "slots":
            items.append(Paragraph("□ " * 8, track_s))

        elif ftype == "image":
            img_data = [[Paragraph("[ Retrato ]", blank_s)]]
            img_tbl  = Table(img_data, colWidths=[3.5*cm], rowHeights=[4*cm])
            img_tbl.setStyle(TableStyle([
                ("BOX",        (0, 0), (-1, -1), 1, C_BORDER),
                ("BACKGROUND", (0, 0), (-1, -1), C_HEADER),
                ("ALIGN",      (0, 0), (-1, -1), "CENTER"),
                ("VALIGN",     (0, 0), (-1, -1), "MIDDLE"),
            ]))
            items.append(img_tbl)

        elif ftype == "checkbox":
            items.append(Paragraph("☐  Sim     ☐  Não", track_s))

        else:
            items.append(Paragraph("_" * 35, blank_s))

        return items

    # ── Renderizar campos ─────────────────────────────────────────────────────
    CELL_STYLE = TableStyle([
        ("BOX",       (0, 0), (-1, -1), 0.8, C_BORDER),
        ("INNERGRID", (0, 0), (-1, -1), 0.4, C_BORDER),
        ("BACKGROUND",(0, 0), (-1, -1), C_WHITE),
        ("VALIGN",    (0, 0), (-1, -1), "TOP"),
        ("PADDING",   (0, 0), (-1, -1), 7),
    ])
    CELL_STYLE_WIDE = TableStyle([
        ("BOX",       (0, 0), (-1, -1), 0.8, C_BORDER),
        ("BACKGROUND",(0, 0), (-1, -1), C_WHITE),
        ("VALIGN",    (0, 0), (-1, -1), "TOP"),
        ("PADDING",   (0, 0), (-1, -1), 7),
    ])

    # Agrupar campos em seções se possível
    current_section = None
    section_story   = []

    def _flush_section():
        if section_story:
            story.extend(section_story)
            section_story.clear()

    i = 0
    while i < len(fields):
        f      = fields[i]
        fwidth = f.get("width", 1)
        ftype  = f.get("type", "text")
        flabel = f.get("label", "")

        is_wide = fwidth == 2 or ftype in ("text_long", "table", "list", "image")

        if is_wide:
            cell_content = _make_cell(flabel, ftype, f)
            tbl = Table([[cell_content]], colWidths=[FULL_W])
            tbl.setStyle(CELL_STYLE_WIDE)
            story.append(tbl)
            story.append(Spacer(1, 0.18*cm))
            i += 1
        else:
            right_content = [Paragraph("", blank_s)]
            if i + 1 < len(fields) and fields[i + 1].get("width", 1) == 1 \
                    and fields[i + 1].get("type") not in ("text_long", "table", "list", "image"):
                f2 = fields[i + 1]
                right_content = _make_cell(f2.get("label", ""), f2.get("type", "text"), f2)
                i += 1

            tbl = Table([[_make_cell(flabel, ftype, f), right_content]], colWidths=COL_W)
            tbl.setStyle(CELL_STYLE)
            story.append(tbl)
            story.append(Spacer(1, 0.18*cm))
            i += 1

    # ── Seção de habilidades especiais ────────────────────────────────────────
    key_abilities = answers.get("key_abilities") or []
    if key_abilities:
        story.append(Spacer(1, 0.2*cm))
        story.append(HRFlowable(width="100%", thickness=1, color=C_BORDER))
        story.append(Paragraph("HABILIDADES ESPECIAIS", sec_s))
        ab_rows = []
        row = []
        for ab in key_abilities:
            row.append(Paragraph(f"◆  {ab}", body_s))
            if len(row) == 2:
                ab_rows.append(row)
                row = []
        if row:
            row.append(Paragraph("", body_s))
            ab_rows.append(row)
        if ab_rows:
            ab_tbl = Table(ab_rows, colWidths=COL_W)
            ab_tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), C_HEADER),
                ("BOX",        (0, 0), (-1, -1), 0.5, C_BORDER),
                ("INNERGRID",  (0, 0), (-1, -1), 0.3, C_BORDER),
                ("PADDING",    (0, 0), (-1, -1), 5),
            ]))
            story.append(ab_tbl)

    # ── Seção de itens / equipamento ──────────────────────────────────────────
    eq_cats = answers.get("equipment_categories") or []
    if eq_cats and "sem_itens" not in eq_cats:
        story.append(Spacer(1, 0.2*cm))
        story.append(HRFlowable(width="100%", thickness=1, color=C_BORDER))
        story.append(Paragraph("EQUIPAMENTOS & ITENS", sec_s))
        eq_header = [Paragraph("Item / Descrição", lbl_s),
                     Paragraph("Peso / Custo", lbl_s),
                     Paragraph("Efeito / Notas", lbl_s)]
        eq_rows = [eq_header] + [[Paragraph("", blank_s)] * 3 for _ in range(8)]
        eq_tbl  = Table(eq_rows, colWidths=[7*cm, 3.5*cm, 7.1*cm])
        eq_tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0),  C_LIGHT),
            ("BOX",        (0, 0), (-1, -1), 0.8, C_BORDER),
            ("INNERGRID",  (0, 0), (-1, -1), 0.4, C_BORDER),
            ("PADDING",    (0, 0), (-1, -1), 5),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [C_WHITE, C_HEADER]),
        ]))
        story.append(eq_tbl)

    # ── Notas finais ──────────────────────────────────────────────────────────
    story.append(Spacer(1, 0.3*cm))
    story.append(HRFlowable(width="100%", thickness=1, color=C_BORDER))
    notes_data = [[
        Paragraph("NOTAS / HISTÓRIA", lbl_s),
    ]]
    for _ in range(5):
        notes_data.append([Paragraph("_" * 85, blank_s)])
    notes_tbl = Table(notes_data, colWidths=[FULL_W])
    notes_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0),  C_LIGHT),
        ("BACKGROUND", (0, 1), (-1, -1), C_WHITE),
        ("BOX",        (0, 0), (-1, -1), 0.8, C_BORDER),
        ("PADDING",    (0, 0), (-1, -1), 6),
    ]))
    story.append(notes_tbl)

    # ── Footer ────────────────────────────────────────────────────────────────
    story.append(Spacer(1, 0.3*cm))
    footer_data = [[
        Paragraph(f"<b>{name}</b>  •  Base: {base}  •  Gerado por DM - Dungeon Music  •  "
                  f"{datetime.now().strftime('%d/%m/%Y')}", sub_s),
    ]]
    footer_tbl = Table(footer_data, colWidths=[FULL_W])
    footer_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), C_PRIMARY),
        ("PADDING",    (0, 0), (-1, -1), 5),
    ]))
    story.append(footer_tbl)

    doc.build(story)
    return output_path


# ─────────────────────────────────────────────────────────────────────────────
# API PÚBLICA
# ─────────────────────────────────────────────────────────────────────────────

def export(system_data: dict, system_formats: list, sheet_formats: list,
           base_dir: str = None) -> dict:
    """
    Exporta documento do sistema e ficha de personagem.

    system_formats: ["txt", "word", "png", "pdf"]  — para o doc do sistema
    sheet_formats:  ["txt", "word", "png", "pdf"]  — para a ficha

    Retorna dict {"sistema_txt": path, "ficha_pdf": path, ...}
    """
    _ensure_dir()
    out_dir = base_dir or _EXPORT_DIR
    os.makedirs(out_dir, exist_ok=True)

    name = system_data.get("name", "sistema")
    safe = _safe_filename(name)
    ts   = datetime.now().strftime("%Y%m%d_%H%M")
    results = {}

    _SYSTEM_FUNCS = {
        "txt":  (export_system_txt,  f"{safe}_sistema_{ts}.txt"),
        "word": (export_system_word, f"{safe}_sistema_{ts}.docx"),
        "png":  (export_system_png,  f"{safe}_sistema_{ts}.png"),
        "pdf":  (export_system_pdf,  f"{safe}_sistema_{ts}.pdf"),
    }
    _SHEET_FUNCS = {
        "txt":  (export_sheet_txt,  f"{safe}_ficha_{ts}.txt"),
        "word": (export_sheet_word, f"{safe}_ficha_{ts}.docx"),
        "png":  (export_sheet_png,  f"{safe}_ficha_{ts}.png"),
        "pdf":  (export_sheet_pdf,  f"{safe}_ficha_{ts}.pdf"),
    }

    for fmt in system_formats:
        if fmt not in _SYSTEM_FUNCS:
            continue
        func, fname = _SYSTEM_FUNCS[fmt]
        key = f"sistema_{fmt}"
        try:
            results[key] = func(system_data, os.path.join(out_dir, fname))
        except Exception as e:
            results[key] = f"ERRO: {e}"

    for fmt in sheet_formats:
        if fmt not in _SHEET_FUNCS:
            continue
        func, fname = _SHEET_FUNCS[fmt]
        key = f"ficha_{fmt}"
        try:
            results[key] = func(system_data, os.path.join(out_dir, fname))
        except Exception as e:
            results[key] = f"ERRO: {e}"

    return results
